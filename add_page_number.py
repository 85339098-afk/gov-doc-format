"""
页码插入 — 公文格式规范

插入 PAGE 域代码实现自动编号。
格式：— N —（N 两侧各一字线）
奇数页居右，偶数页居左，首页不显示。

域代码插入方式说明：
- 操作 run._r（私有属性）是 python-docx 生态中插入域代码的通行做法
- Word 的域代码规范要求 fldChar/instrText 必须嵌在 w:r 元素内部
- 本实现为每个域代码片段创建独立 run，减少单个 run 的复杂度
"""

from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from .set_font import set_run_font


def _add_field_run(paragraph, field_type, text=None):
    """
    在段落中添加一个域代码专用的 run。

    创建一个独立的 run，在其 _r 元素中插入域代码组件。
    每个域代码片段（begin / instrText / end）占用独立 run，
    避免把多个元素塞进同一个 run 的 _r 中。

    Args:
        paragraph: 目标段落
        field_type: 'begin' | 'instr' | 'end'
        text: instrText 的文本内容（仅 field_type='instr' 时需要）
    """
    run = paragraph.add_run()
    set_run_font(run, '宋体', 14, bold=False)

    if field_type in ('begin', 'end'):
        fld = OxmlElement('w:fldChar')
        fld.set(qn('w:fldCharType'), field_type)
        run.element.append(fld)
    elif field_type == 'instr' and text:
        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = text
        run.element.append(instr)


def _insert_page_number_run(paragraph):
    """
    在段落中插入 "— PAGE —" 域代码。

    分 5 个独立 run 插入，确保 Word 正确渲染：
    run1: "— "（文本）
    run2: fldChar begin
    run3: instrText PAGE
    run4: fldChar end
    run5: " —"（文本）
    """
    # run1: "— " 文本
    run = paragraph.add_run('— ')
    set_run_font(run, '宋体', size=Pt(14), bold=False)

    # run2: fldChar begin
    _add_field_run(paragraph, 'begin')

    # run3: instrText PAGE
    _add_field_run(paragraph, 'instr', ' PAGE ')

    # run4: fldChar end
    _add_field_run(paragraph, 'end')

    # run5: " —" 文本
    run = paragraph.add_run(' —')
    set_run_font(run, '宋体', size=Pt(14), bold=False)


def add_page_number(section):
    """
    为文档节添加页码。

    奇数页页码居右，偶数页页码居左。
    首页不显示页码（需设置 different_first_page_header_footer=True）。

    Args:
        section: Document section 对象
    """
    # === 奇数页页码（居右） ===
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _insert_page_number_run(p)

    # === 偶数页页码（居左） ===
    even_footer = section.even_page_footer
    even_footer.is_linked_to_previous = False
    p = even_footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _insert_page_number_run(p)

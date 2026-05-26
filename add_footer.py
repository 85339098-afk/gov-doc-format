"""
版记生成 — 公文格式规范

生成公文版记，包括抄送机关、印发机关、印发日期。
位于公文最后一页底部。

版记中"印发日期右空一字"的实现：
- 使用 \t（制表符）+ 右对齐制表位，将印发日期推到行末减一字宽
- 右空一字 = 从页面右边距向左缩进一个中文字宽 ≈ 16pt
"""
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from .set_font import set_run_font, set_paragraph_spacing


def add_cc_line(doc, text):
    """
    添加抄送机关。

    格式：抄送机关后加句号

    Args:
        doc: Document 对象
        text: 抄送内容（如 "抄送：市总工会，市财政局。"）
    """
    p = doc.add_paragraph()
    set_paragraph_spacing(p, line_spacing=Pt(28), space_before=Pt(0), space_after=Pt(0))

    run = p.add_run(text)
    set_run_font(run, '仿宋_GB2312', size=Pt(14), bold=False)  # GB/T 9704-2012：版记用 4号(14pt) 仿宋

    return p


def add_dispatch_info(doc, office="", date_str=""):
    """
    添加印发机关和印发日期。

    格式：印发机关左对齐，印发日期右空一字（使用制表位实现精确右对齐）

    Args:
        doc: Document 对象
        office: 印发机关名称
        date_str: 印发日期
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    set_paragraph_spacing(p, line_spacing=Pt(28), space_before=Pt(0), space_after=Pt(0))

    # 设置右对齐制表位：从右边界回退一个字宽（16pt）
    # 动态计算：从文档节获取版心宽度，避免硬编码 ≈ 442pt，右空一字 ≈ 从右缩进一个字的宽度
    section = doc.sections[0]
    ruler_w = section.page_width - section.left_margin - section.right_margin - Pt(16)
    pf.tab_stops.add_tab_stop(ruler_w, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)


    # 印发机关（左对齐）+ 制表符 + 印发日期（右对齐）
    run = p.add_run(office)
    set_run_font(run, '仿宋_GB2312', size=Pt(14))

    run = p.add_run("\t" + date_str)
    set_run_font(run, '仿宋_GB2312', size=Pt(14))

    return p

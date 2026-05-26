"""
落款生成 — 公文格式规范

生成发文字号、签发人、落款单位名称和成文日期。
"""

from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from .set_font import set_run_font, set_paragraph_spacing


def add_doc_number(doc, text):
    """
    添加发文字号。

    格式：机关代字〔年份〕顺序号
    示例：工办〔2026〕15号
    要求：标题下空一行，居中排列

    Args:
        doc: Document 对象
        text: 发文字号文本（如 "工办〔2026〕15号"）
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, line_spacing=Pt(28), space_before=Pt(0), space_after=Pt(0))

    run = p.add_run(text)
    set_run_font(run, '仿宋_GB2312', size=Pt(16), bold=False)

    return p


def add_signee(doc, text):
    """
    添加签发人（仅上行文（请示、报告）使用）。

    格式：与发文字号同行，签发人居右。
    实际使用中请将发文字号和签发人放在同一行。

    Args:
        doc: Document 对象
        text: 签发人文本
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(p, line_spacing=Pt(28), space_before=Pt(0), space_after=Pt(0))

    # "签发人" 三字用 3 号仿宋
    run = p.add_run("签发人：")
    set_run_font(run, '仿宋_GB2312', size=Pt(16), bold=False)

    # 签发人姓名用 3 号楷体
    run = p.add_run(text.replace("签发人：", ""))
    set_run_font(run, '楷体_GB2312', size=Pt(16), bold=False)

    return p


def add_organization(doc, name):
    """
    添加落款单位名称。

    格式：成文日期上方居中，与日期之间空一行
    单位名称右空二字排布（加盖公章时使用）

    Args:
        doc: Document 对象
        name: 单位名称
    """
    # 先添加一个空行
    p_space = doc.add_paragraph()
    set_paragraph_spacing(p_space, line_spacing=Pt(28), space_before=Pt(0), space_after=Pt(0))

    # 单位名称（右空二字 = 右缩进 2字符 = 32pt）
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pf = p.paragraph_format
    pf.line_spacing = Pt(28)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Pt(0)
    pf.right_indent = Pt(32)  # 右空二字 ≈ 32pt

    run = p.add_run(name)
    set_run_font(run, '仿宋_GB2312', size=Pt(16), bold=False)

    return p


def add_date(doc, date_str="2026年5月25日"):
    """
    添加成文日期。

    格式：右空四字（右缩进 4 字符 ≈ 64pt），用阿拉伯数字
    含 right_indent 实现右空四字，配合右对齐精确控制落款位置。

    Args:
        doc: Document 对象
        date_str: 日期字符串
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pf = p.paragraph_format
    pf.line_spacing = Pt(28)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Pt(0)
    pf.right_indent = Pt(64)  # 右空四字 ≈ 4字符 × 16pt

    run = p.add_run(date_str)
    set_run_font(run, '仿宋_GB2312', size=Pt(16), bold=False)

    return p

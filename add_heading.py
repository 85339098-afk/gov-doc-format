"""
标题生成 — 公文格式规范

提供各级标题和公文标题的生成函数，符合 GB/T 9704-2012 排版要求。
"""

from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .set_font import set_run_font, set_paragraph_spacing, set_first_line_indent


def add_title(doc, text):
    """
    添加公文标题。

    格式：方正小标宋简体 二号(22pt) 居中

    Args:
        doc: Document 对象
        text: 标题文本

    Returns:
        Paragraph 对象
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, line_spacing=Pt(28), space_before=Pt(0), space_after=Pt(0))

    run = p.add_run(text)
    set_run_font(run, '方正小标宋简体', size=Pt(22), bold=False)

    return p


def add_heading1(doc, text):
    """
    添加一级标题（一、二、三……）。

    格式：黑体 三号(16pt) 左空二字（不加粗）
    统一使用 add_paragraph，不使用 Word 内置标题样式，避免样式污染。

    Args:
        doc: Document 对象
        text: 标题文本

    Returns:
        Paragraph 对象
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, line_spacing=Pt(28), space_before=Pt(0), space_after=Pt(0))
    set_first_line_indent(p)

    run = p.add_run(text)
    set_run_font(run, '黑体', size=Pt(16), bold=False)

    return p


def add_heading2(doc, text, bold=False):
    """
    添加二级标题（（一）（二）（三）……）。

    格式：楷体 三号(16pt) 左空二字

    Args:
        doc: Document 对象
        text: 标题文本
        bold: 是否加粗，默认 False

    Returns:
        Paragraph 对象
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, line_spacing=Pt(28), space_before=Pt(0), space_after=Pt(0))
    set_first_line_indent(p)

    run = p.add_run(text)
    set_run_font(run, '楷体_GB2312', size=Pt(16), bold=bold)

    return p


def add_heading3(doc, text, bold=False):
    """
    添加三级标题（1. 2. 3.……）。

    格式：仿宋 三号(16pt) 左空二字，可加粗
    与正文区别：三级标题行距压缩，有独立的语义标记。

    Args:
        doc: Document 对象
        text: 标题文本（含编号，如 "1. 字体要求"）
        bold: 是否加粗，默认 False

    Returns:
        Paragraph 对象
    """
    p = doc.add_paragraph()
    set_paragraph_spacing(p, line_spacing=Pt(28), space_before=Pt(0), space_after=Pt(0))
    set_first_line_indent(p)

    run = p.add_run(text)
    set_run_font(run, '仿宋_GB2312', size=Pt(16), bold=bold)

    return p

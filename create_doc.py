"""
主入口脚本 — 公文格式规范

创建标准公文文档，整合所有子模块：
1. 创建 A4 文档，设置页边距
2. 设置奇偶页和首页不同的页眉页脚
3. 插入页码
"""

from docx import Document
from docx.shared import Cm

from .add_heading import add_title, add_heading1, add_heading2, add_heading3
from .add_page_number import add_page_number
from .format_body import add_body_text
from .add_signature import add_doc_number, add_signee, add_organization, add_date
from .add_footer import add_cc_line, add_dispatch_info


def create_official_doc():
    """
    创建标准公文文档。

    返回:
        tuple: (Document, Section)
    """
    doc = Document()

    # ============================================================
    # 页面设置
    # ============================================================
    section = doc.sections[0]
    section.page_width = Cm(21.0)       # A4
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.7)        # 上 37mm
    section.bottom_margin = Cm(3.5)     # 下 35mm
    section.left_margin = Cm(2.8)       # 左 28mm
    section.right_margin = Cm(2.6)      # 右 26mm

    # ============================================================
    # 页眉页脚设置
    # ============================================================
    # 首页不显示页码
    section.different_first_page_header_footer = True   # titlePg
    # 奇偶页页码分别设置
    section.even_and_odd_header_footer = True            # evenAndOddHeaders

    # ============================================================
    # 插入页码
    # ============================================================
    add_page_number(section)

    return doc, section


# ============================================================
# 使用示例
# ============================================================
if __name__ == "__main__":
    doc, section = create_official_doc()

    # 公文标题
    add_title(doc, "关于印发《公文格式规范》的通知")

    # 发文字号
    add_doc_number(doc, "工办〔2026〕15号")

    # 主送机关（手动添加）
    add_body_text(doc, "各科室、各下属单位：")

    # 一级标题
    add_heading1(doc, "一、总体要求")

    # 正文
    add_body_text(doc, "为规范公文格式，提高办文质量，根据《党政机关公文处理工作条例》及GB/T 9704-2012《党政机关公文格式》有关规定，结合实际，制定本规范。")

    # 二级标题
    add_heading2(doc, "（一）页面设置")

    # 正文
    add_body_text(doc, "公文用纸采用A4型纸（210mm×297mm），页边距为上37mm、下35mm、左28mm、右26mm，版心尺寸为156mm×225mm。左侧装订。")

    # 三级标题
    add_heading3(doc, "1. 字体字号要求")

    add_body_text(doc, "公文标题使用2号方正小标宋简体，一级标题使用3号黑体，二级标题使用3号楷体，正文使用3号仿宋。")

    # 落款
    add_organization(doc, "×××办公室")
    add_date(doc, "2026年5月25日")

    # 版记
    add_cc_line(doc, "抄送：×××，×××。")
    add_dispatch_info(doc, "×××办公室", "2026年5月25日印发")

    doc.save("正式公文_示例.docx")
    print("✅ 公文文档已生成：正式公文_示例.docx")

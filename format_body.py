"""
正文排版 — 公文格式规范

设置正文段落格式：
- 行距固定值 28 磅
- 首行缩进 2 字符
- 段前段后 0
- 字体仿宋三号
"""

from docx.shared import Pt
from .set_font import set_run_font, set_paragraph_spacing, set_first_line_indent


def add_body_text(doc, text, font_name='仿宋_GB2312', size=Pt(16)):
    """
    添加正文段落。

    Args:
        doc: Document 对象
        text: 正文文本
        font_name: 字体名称，默认仿宋_GB2312
        size: 字号，默认 16pt（三号）

    Returns:
        Paragraph 对象
    """
    LINE_SPACING = Pt(28)  # 多数机关沿用 28pt；严格对标标准 ≈ 29.5pt (GB/T 33476.2-2016)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, line_spacing=LINE_SPACING, space_before=Pt(0), space_after=Pt(0))
    set_first_line_indent(p)

    run = p.add_run(text)
    set_run_font(run, font_name, size=size, bold=False)

    return p


def add_body_text_no_indent(doc, text, font_name='仿宋_GB2312', size=Pt(16)):
    """
    添加无缩进的正文段落（用于特殊情况，如标题换行后的续行）。

    Args:
        doc: Document 对象
        text: 正文文本
        font_name: 字体名称
        size: 字号
    """
    p = doc.add_paragraph()
    set_paragraph_spacing(p, line_spacing=Pt(28), space_before=Pt(0), space_after=Pt(0))

    run = p.add_run(text)
    set_run_font(run, font_name, size=size, bold=False)

    return p
